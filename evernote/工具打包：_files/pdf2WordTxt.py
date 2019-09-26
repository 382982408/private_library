#!/usr/bin/python
# -*- coding:utf-8 -*-

###lib required:pdfminer3k,python_docx
import sys,io,re,os
from pdfminer.pdfinterp import PDFResourceManager, process_pdf
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from io import StringIO
from io import open
from docx import Document
from configparser import ConfigParser
from concurrent.futures import ProcessPoolExecutor
sys.stdout = io.TextIOWrapper(sys.stdout.buffer,encoding='utf-8')

def read_from_pdf(file_path):
    with open(file_path,'rb') as file:
        resource_manager=PDFResourceManager()
        return_str=StringIO()
        lap_params=LAParams()
        device=TextConverter(resource_manager,return_str,laparams=lap_params)
        process_pdf(resource_manager,device,file)
        device.close()
        content=return_str.getvalue()
        return_str.close()
        return content

def readPDF(pdfFile):
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, laparams=laparams) 
    process_pdf(rsrcmgr, device, pdfFile)
    device.close()
    content = retstr.getvalue()
    retstr.close()
    return content

def save2txt(out_file,txt_content):
    with open(out_file,"w",encoding='utf-8') as f:
        f.write(txt_content)

def save_text_to_word(content,file_path):
    doc = Document()
    for line in content.split('\n'):
        paragraph=doc.add_paragraph()
        paragraph.add_run(remove_control_characters(line))
    doc.save(file_path)
 
def remove_control_characters(content):
    mpa=dict.fromkeys(range(32))
    return content.translate(mpa)
 
def pdf_to_word(pdf_file_path,word_file_path):
    content=read_from_pdf(pdf_file_path)
    save_text_to_word(content,word_file_path)

def main():
    if len(sys.argv)<2:
        print (f"\n\tpython {sys.argv[0]} <input_pdf>\n")
        sys.exit(0)
    else:
        txt_out_file=sys.argv[1].replace("pdf","txt")
        docx_out_file=sys.argv[1].replace("pdf","docx")

        ####pdf转成txt
        pdf_content = readPDF(open(sys.argv[1], 'rb'))
        save2txt(txt_out_file,pdf_content)

        ####pdf转成word
        with ProcessPoolExecutor(max_workers=1) as executor:
            result=executor.submit(pdf_to_word,sys.argv[1],docx_out_file)
 
if __name__=='__main__':
    main()
